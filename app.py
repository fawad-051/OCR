import streamlit as st
from PIL import Image
import io
import os
import pytesseract
import platform

if platform.system() == "Windows":
    pytesseract.pytesseract.tesseract_cmd = r"C:\Program Files\Tesseract-OCR\tesseract.exe"
else:
    pytesseract.pytesseract.tesseract_cmd = "/usr/bin/tesseract"

import tempfile
import cv2
import numpy as np
from pdf2image import convert_from_bytes
import pandas as pd
from docx import Document
from docx.shared import Pt

# ---------- Configuration (modify paths if needed) ----------
# If tesseract is not in PATH, set the pytesseract.tesseract_cmd to the full executable path, e.g.:
# pytesseract.pytesseract.tesseract_cmd = r"C:\Program Files\Tesseract-OCR\tesseract.exe"

# ---------- Helper functions ----------

def load_images_from_upload(uploaded_file):
    """Return list of PIL Images for given uploaded file (pdf or image)."""
    file_bytes = uploaded_file.read()
    uploaded_file.seek(0)
    if uploaded_file.type == "application/pdf" or uploaded_file.name.lower().endswith('.pdf'):
        # convert pdf to images (one per page)
        images = convert_from_bytes(file_bytes)
    else:
        image = Image.open(io.BytesIO(file_bytes)).convert("RGB")
        images = [image]
    return images


def ocr_image_to_text(image: Image.Image, lang='eng'):
    """Return plain text from image using pytesseract."""
    return pytesseract.image_to_string(image, lang=lang)


def detect_table_and_extract(image: Image.Image, lang='eng'):
    """Try to detect table structure and extract a pandas DataFrame. If no table found, return None.

    The approach:
    - Convert to grayscale
    - Adaptive threshold
    - Use morphological operations to detect horizontal and vertical lines
    - Find intersections to detect grid, then segment cells
    - OCR each cell
    """
    # Convert PIL -> OpenCV image
    cv_img = cv2.cvtColor(np.array(image), cv2.COLOR_RGB2BGR)
    gray = cv2.cvtColor(cv_img, cv2.COLOR_BGR2GRAY)
    # adaptive threshold to binary
    thr = cv2.adaptiveThreshold(~gray,255,cv2.ADAPTIVE_THRESH_MEAN_C,cv2.THRESH_BINARY,15,-2)

    # Detect horizontal lines
    horizontal = thr.copy()
    cols = horizontal.shape[1]
    horizontal_size = max(10, cols // 30)
    horizontalStructure = cv2.getStructuringElement(cv2.MORPH_RECT, (horizontal_size,1))
    horizontal = cv2.erode(horizontal, horizontalStructure)
    horizontal = cv2.dilate(horizontal, horizontalStructure)

    # Detect vertical lines
    vertical = thr.copy()
    rows = vertical.shape[0]
    vertical_size = max(10, rows // 30)
    verticalStructure = cv2.getStructuringElement(cv2.MORPH_RECT, (1, vertical_size))
    vertical = cv2.erode(vertical, verticalStructure)
    vertical = cv2.dilate(vertical, verticalStructure)

    # Combine lines
    mask = horizontal + vertical

    # Find joints (intersection points)
    joints = cv2.bitwise_and(horizontal, vertical)

    # Find contours from mask to locate table regions
    contours, hierarchy = cv2.findContours(mask, cv2.RETR_EXTERNAL, cv2.CHAIN_APPROX_SIMPLE)
    table_contours = []
    for cnt in contours:
        x,y,w,h = cv2.boundingRect(cnt)
        if w > 50 and h > 50:  # filter small regions
            table_contours.append((x,y,w,h))

    if not table_contours:
        return None

    # Choose largest contour (assume largest table on page)
    table_contours = sorted(table_contours, key=lambda x: x[2]*x[3], reverse=True)
    x,y,w,h = table_contours[0]
    table_roi = ~thr[y:y+h, x:x+w]
    hor = horizontal[y:y+h, x:x+w]
    ver = vertical[y:y+h, x:x+w]
    joint = joints[y:y+h, x:x+w]

    # Find cell corners via joint points
    corners = cv2.findNonZero(joint)
    if corners is None or len(corners) < 4:
        # not enough grid intersections
        return None

    # Project points to grid by clustering x and y coordinates
    points = [(pt[0][0], pt[0][1]) for pt in corners]
    xs = sorted([p[0] for p in points])
    ys = sorted([p[1] for p in points])

    # cluster close coordinates to get unique grid lines
    def cluster_coords(coords, eps=10):
        clusters = []
        current = [coords[0]]
        for c in coords[1:]:
            if abs(c - current[-1]) <= eps:
                current.append(c)
            else:
                clusters.append(int(sum(current)/len(current)))
                current = [c]
        clusters.append(int(sum(current)/len(current)))
        return clusters

    xs_c = cluster_coords(xs, eps=max(5, w//100))
    ys_c = cluster_coords(ys, eps=max(5, h//100))

    # number of rows and cols
    ncols = len(xs_c)-1
    nrows = len(ys_c)-1
    if ncols <=0 or nrows <= 0 or ncols>50 or nrows>200:
        return None

    # Build table by extracting each cell
    table = []
    for r in range(nrows):
        row_cells = []
        for c in range(ncols):
            x1 = xs_c[c]
            x2 = xs_c[c+1]
            y1 = ys_c[r]
            y2 = ys_c[r+1]
            cell = table_roi[y1:y2, x1:x2]
            # pad small region
            if cell.size == 0:
                text = ''
            else:
                # OCR the cell - convert to PIL
                cell_rgb = cv2.cvtColor(cell, cv2.COLOR_GRAY2RGB)
                pil_cell = Image.fromarray(cell_rgb)
                text = pytesseract.image_to_string(pil_cell, lang=lang)
                text = text.replace('\n', ' ').strip()
            row_cells.append(text)
        table.append(row_cells)

    df = pd.DataFrame(table)
    return df


def image_to_docx(images, output_docx_path, lang='eng'):
    """Process images: try to extract tables; if not, extract page text; write to docx."""
    doc = Document()
    # small default style
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)

    for i, img in enumerate(images, start=1):
        st.write(f"Processing page {i} ...")
        df = detect_table_and_extract(img, lang=lang)
        if isinstance(df, pd.DataFrame):
            # write table to docx
            rows, cols = df.shape
            doc.add_paragraph(f"Table detected on page {i} (rows={rows}, cols={cols})")
            table = doc.add_table(rows+1, cols)
            # header as first row from df.iloc[0] if it looks like header
            for c in range(cols):
                table.cell(0,c).text = str(df.iloc[0,c])
            for r in range(1, rows):
                for c in range(cols):
                    table.cell(r, c).text = str(df.iloc[r,c])
            doc.add_paragraph('\n')
        else:
            # fallback to full-page OCR
            text = ocr_image_to_text(img, lang=lang)
            doc.add_paragraph(text)
            doc.add_page_break()

    doc.save(output_docx_path)
    return output_docx_path


# ---------- Streamlit UI ----------

def main():
    st.set_page_config(page_title="OCR → Word Converter", layout="centered")
    st.title("Scans and transcribes your documents")
    st.write("Upload an image or PDF. The app will OCR the content, attempt to detect tables and extract them as tables in the Word file. If table detection fails, full-page text will be exported.")

    uploaded_file = uploaded_file = st.file_uploader(
    "Upload image or PDF", 
    type=['pdf','png','jpg','jpeg','tiff','tif'], 
    accept_multiple_files=False,
    label_visibility="visible"
)

    if uploaded_file is not None:
        with st.spinner('Loading and converting file...'):
            try:
                images = load_images_from_upload(uploaded_file)
            except Exception as e:
                st.error(f"Failed to load file: {e}")
                return

        st.write(f"Pages found: {len(images)}")

        if st.button('Run OCR & Convert to Word'):
            with st.spinner('Running OCR — this may take a while for large files...'):
                try:
                    out_path = os.path.join(tempfile.gettempdir(), f"ocr_output_{uploaded_file.name}.docx")
                    image_to_docx(images, out_path, lang=lang)
                except Exception as e:
                    st.error(f"Error while processing: {e}")
                    return

            st.success('Conversion complete — download below')
            with open(out_path, 'rb') as f:
                btn = st.download_button(
                    label='Download .docx',
                    data=f,
                    file_name=f"{os.path.splitext(uploaded_file.name)[0]}.docx",
                    mime='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
                )

    st.markdown("---")
    st.markdown("**Notes / Requirements:**")
    st.markdown("- Install Tesseract OCR (separate from Python). On Windows: install Tesseract and set `pytesseract.pytesseract.tesseract_cmd` if not in PATH. On Linux: `sudo apt install tesseract-ocr`.")
    st.markdown("- For PDF support, install `poppler` (for pdf2image). On Ubuntu: `sudo apt install poppler-utils`. On Windows: install Poppler and add to PATH.")
    st.markdown("- Python libraries: `pip install streamlit pillow pytesseract pdf2image opencv-python-headless python-docx pandas`")
    st.markdown("- Table extraction tries to detect grid-like tables. It works best on clear scanned documents with visible cell borders. Complex or borderless tables may not be perfectly reconstructed; in those cases the page text is exported instead.")


if __name__ == '__main__':
    main()


